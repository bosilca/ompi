#!/usr/bin/env python3
"""Generate a Word document from the k_allreduce design documents."""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))

def set_cell_shading(cell, color):
    """Set background color on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    tcPr.append(shading)

def add_code_block(doc, text):
    """Add a formatted code block."""
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_table_from_rows(doc, header, rows):
    """Add a formatted table."""
    ncols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = h.strip()
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = val.strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

def clean_latex(text):
    """Convert LaTeX-ish notation to readable Unicode/plaintext."""
    text = re.sub(r'\\\[', '', text)
    text = re.sub(r'\\\]', '', text)
    text = re.sub(r'\\\(', '', text)
    text = re.sub(r'\\\)', '', text)
    text = re.sub(r'\\boxed\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\underbrace\{([^}]+)\}_\{([^}]+)\}', r'\1 [\2]', text)
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', text)
    text = re.sub(r'\\left', '', text)
    text = re.sub(r'\\right', '', text)
    text = re.sub(r'\\cdot', '·', text)
    text = re.sub(r'\\times', '×', text)
    text = re.sub(r'\\approx', '≈', text)
    text = re.sub(r'\\geq', '≥', text)
    text = re.sub(r'\\leq', '≤', text)
    text = re.sub(r'\\neq', '≠', text)
    text = re.sub(r'\\gg', '≫', text)
    text = re.sub(r'\\ll', '≪', text)
    text = re.sub(r'\\alpha', 'α', text)
    text = re.sub(r'\\beta', 'β', text)
    text = re.sub(r'\\gamma', 'γ', text)
    text = re.sub(r'\\ell_k', 'ℓ_k', text)
    text = re.sub(r'\\ell_G', 'ℓ_G', text)
    text = re.sub(r'\\ell_N', 'ℓ_N', text)
    text = re.sub(r'\\lceil', '⌈', text)
    text = re.sub(r'\\rceil', '⌉', text)
    text = re.sub(r'\\lfloor', '⌊', text)
    text = re.sub(r'\\rfloor', '⌋', text)
    text = re.sub(r'\\log_2', 'log₂', text)
    text = re.sub(r'\\ln', 'ln', text)
    text = re.sub(r'\\sum_\{([^}]*)\}\^\{([^}]*)\}', r'Σ(\1 to \2)', text)
    text = re.sub(r'\\mid', '|', text)
    text = re.sub(r'\\nmid', '∤', text)
    text = re.sub(r'\\sim', '~', text)
    text = re.sub(r'\$', '', text)
    text = re.sub(r'\\,', ' ', text)
    text = re.sub(r'\\;', ' ', text)
    text = re.sub(r'\\quad', '  ', text)
    text = re.sub(r'\\partial', '∂', text)
    text = text.replace('\\\\', '\n')
    return text.strip()

def parse_md_table(lines, start):
    """Parse a markdown table starting at line index start. Returns (header, rows, end_idx)."""
    header_line = lines[start].strip()
    cols = [c.strip() for c in header_line.split('|')[1:-1]]

    rows = []
    i = start + 2
    while i < len(lines) and lines[i].strip().startswith('|'):
        row = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
        rows.append(row)
        i += 1
    return cols, rows, i

def process_md(doc, md_text, title_prefix=""):
    """Process markdown text and add it to the Word document."""
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        # Code block boundaries
        if line.strip().startswith('```'):
            if in_code:
                add_code_block(doc, '\n'.join(code_buf))
                code_buf = []
                in_code = False
                i += 1
                continue
            else:
                in_code = True
                i += 1
                continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            doc.add_heading(clean_latex(stripped[2:]), level=1)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(clean_latex(stripped[3:]), level=2)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(clean_latex(stripped[4:]), level=3)
            i += 1
            continue
        if stripped.startswith('#### '):
            doc.add_heading(clean_latex(stripped[5:]), level=4)
            i += 1
            continue

        # Tables
        if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            cols, rows, end = parse_md_table(lines, i)
            cols = [clean_latex(c) for c in cols]
            rows = [[clean_latex(c) for c in r] for r in rows]
            add_table_from_rows(doc, cols, rows)
            doc.add_paragraph()
            i = end
            continue

        # LaTeX display math (block)
        if stripped.startswith('\\['):
            math_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('\\]'):
                math_lines.append(lines[i])
                i += 1
            i += 1
            math_text = clean_latex(' '.join(math_lines))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(math_text)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(10)
            run.italic = True
            continue

        # Blockquote
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(clean_latex(stripped[2:]))
            run.italic = True
            i += 1
            continue

        # Bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = clean_latex(stripped[2:])
            # Handle bold at start
            bold_match = re.match(r'\*\*(.+?)\*\*:?\s*(.*)', text)
            p = doc.add_paragraph(style='List Bullet')
            if bold_match:
                run = p.add_run(bold_match.group(1))
                run.bold = True
                rest = bold_match.group(2)
                if rest:
                    p.add_run(': ' + rest if not bold_match.group(0).endswith(':') else rest)
            else:
                # handle inline bold/code
                add_rich_text(p, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            text = clean_latex(num_match.group(2))
            p = doc.add_paragraph(style='List Number')
            add_rich_text(p, text)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph — accumulate continuation lines
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith('#') or nxt.startswith('|') or
                nxt.startswith('```') or nxt.startswith('- ') or
                nxt.startswith('* ') or nxt.startswith('> ') or
                nxt.startswith('\\[') or nxt == '---' or
                re.match(r'^\d+\.\s+', nxt)):
                break
            para_lines.append(nxt)
            i += 1

        text = clean_latex(' '.join(para_lines))
        p = doc.add_paragraph()
        add_rich_text(p, text)

def add_rich_text(paragraph, text):
    """Add text with inline bold (**) and code (`) formatting."""
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        elif part:
            paragraph.add_run(part)


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ── Title page ──
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('2D Grid Allreduce (k_allreduce)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Design, Implementation, and Performance Modeling\n'
                           'for Open MPI')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('ompi_coll_base_allreduce_intra_k_allreduce')
    run.font.name = 'Courier New'
    run.font.size = Pt(11)

    doc.add_page_break()

    # ── Table of contents placeholder ──
    doc.add_heading('Table of Contents', level=1)
    p = doc.add_paragraph('Part I:   Algorithm Design')
    p = doc.add_paragraph('Part II:  LogGP Cost Model')
    p = doc.add_paragraph('Part III: Performance Visualization')
    doc.add_page_break()

    # ── Part I: Design ──
    doc.add_heading('Part I: Algorithm Design', level=1)
    doc.add_paragraph()
    with open(os.path.join(BASE, 'K_ALLREDUCE_DESIGN.md')) as f:
        md = f.read()
    # Skip the top-level heading (already in Part title)
    md = re.sub(r'^# .+?\n', '', md, count=1)
    process_md(doc, md)
    doc.add_page_break()

    # ── Part II: Cost Model ──
    doc.add_heading('Part II: LogGP Cost Model', level=1)
    doc.add_paragraph()
    with open(os.path.join(BASE, 'K_ALLREDUCE_COST_MODEL.md')) as f:
        md = f.read()
    md = re.sub(r'^# .+?\n', '', md, count=1)
    process_md(doc, md)
    doc.add_page_break()

    # ── Part III: Performance Plot ──
    doc.add_heading('Part III: Performance Visualization', level=1)
    doc.add_paragraph()

    p = doc.add_paragraph(
        'The following plot shows the effective allreduce bandwidth predicted '
        'by the LogGP model for a DGX-like system: 128 processes (N=128), '
        '8 GPUs per node via NVLink (k=8), 16 nodes connected by '
        '8\u00d7400 Gbps Ethernet (50 GB/s per NIC, 400 GB/s aggregate per '
        'node). Each GPU has a dedicated NIC, and Phase 2\u2019s column-based '
        'communication is inherently rail-optimized\u2014all 8 NICs are '
        'active simultaneously. Message sizes range from 16 KB to 4 MB.'
    )

    img_path = os.path.join(BASE, 'k_allreduce_bandwidth_model.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph('[Plot not found — run plot_k_allreduce_model.py first]')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Key observations:')
    run.bold = True

    findings = [
        ('C3 (RS + RD + Allgather) achieves the highest bandwidth in this range, '
         'benefiting from reduced intra-node volume (reduce-scatter) and low '
         'inter-node latency (log\u2082(N/k) recursive-doubling round-trips).'),
        ('C4 (RS + Ring + Allgather) is bandwidth- and computation-optimal but '
         'suffers from high inter-node latency (2(N/k\u22121) ring steps). It '
         'overtakes C3 only for messages larger than ~23 MB.'),
        ('All k_allreduce variants outperform flat Rabenseifner for large messages '
         'by reducing inter-node data volume by a factor of k.'),
        ('The Phase 1 choice (reduce-scatter vs allreduce) has the largest impact, '
         'as it determines how much data traverses the NVLink interconnect.'),
        ('Phase 2 is inherently rail-optimized: each of the k=8 independent column '
         'streams flows through a dedicated NIC (GPU p \u2194 NIC p), yielding '
         '8\u00d750 = 400 GB/s aggregate inter-node bandwidth with zero NIC '
         'contention. This property is structural and requires no special '
         'process placement.'),
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    # ── Save ──
    out = os.path.join(BASE, 'K_ALLREDUCE_COMPLETE.docx')
    doc.save(out)
    print(f'Saved → {out}')

if __name__ == '__main__':
    main()
